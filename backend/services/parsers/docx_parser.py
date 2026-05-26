import docx

def extract_text_from_docx(file_path: str) -> str:
    """
    Extracts text from a DOCX file, including paragraphs and tables.
    """
    text = ""
    try:
        doc = docx.Document(file_path)
        
        # Process paragraphs
        for para in doc.paragraphs:
            if para.text:
                text += para.text + "\n"
                
        # Process tables (often containing education or layout grids)
        for table in doc.tables:
            for row in table.rows:
                # Remove duplicates in adjacent cells (common in merged docx cells)
                row_cells = []
                for cell in row.cells:
                    cell_txt = cell.text.strip() if cell.text else ""
                    if not row_cells or row_cells[-1] != cell_txt:
                        row_cells.append(cell_txt)
                
                filtered_row = [txt for txt in row_cells if txt]
                if filtered_row:
                    text += " | ".join(filtered_row) + "\n"
    except Exception as e:
        print(f"Error reading DOCX {file_path}: {e}")
    return text
